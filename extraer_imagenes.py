from docx import Document
import os

# Leer el documento de Junio
doc = Document('JAC Junio BITACORA DE ACTIVIDADES.docx')

# Extraer texto y relaciones de imágenes
print("=== ANÁLISIS DEL DOCUMENTO DE JUNIO ===\n")

# Obtener todas las relaciones de imágenes
image_rels = {}
for rel in doc.part.rels.values():
    if "image" in rel.target_ref:
        image_rels[rel.rId] = rel.target_ref

print(f"Total de relaciones de imágenes encontradas: {len(image_rels)}")
print(f"IDs de imágenes: {list(image_rels.keys())}\n")

# Recorrer párrafos y buscar imágenes
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"Párrafo {i}: {para.text[:150]}")
        
        # Buscar imágenes en este párrafo
        for run in para.runs:
            if run._element.xpath('.//pic:pic'):
                print(f"  -> IMAGEN en este párrafo")
                for drawing in run._element.xpath('.//a:blip'):
                    embed = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed:
                        print(f"     Image ID: {embed}")
                        if embed in image_rels:
                            print(f"     Archivo: {image_rels[embed]}")
        print()

# También revisar tablas
for table_idx, table in enumerate(doc.tables):
    print(f"\n=== TABLA {table_idx} ===")
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            if cell.text.strip():
                print(f"Celda [{row_idx}][{cell_idx}]: {cell.text[:100]}")
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run._element.xpath('.//pic:pic'):
                            print(f"  -> IMAGEN en esta celda")
