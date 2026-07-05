from docx import Document
import os
import shutil

# Leer el documento de Junio
doc = Document('JAC Junio BITACORA DE ACTIVIDADES.docx')

print("=== EXTRAYENDO IMÁGENES EN ORDEN CRONOLÓGICO ===\n")

# Crear directorio para imágenes organizadas
output_dir = 'imagenes_organizadas'
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Contador global de imágenes
image_counter = 0
current_bitacora = "Inicio"

# Función para copiar imagen
def copy_image(source_path, dest_name):
    global image_counter
    image_counter += 1
    dest_path = os.path.join(output_dir, dest_name)
    if os.path.exists(source_path):
        shutil.copy(source_path, dest_path)
        print(f"  ✓ Copiada como: {dest_name}")
        return True
    return False

# Recorrer el documento buscando bitácoras e imágenes
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Detectar inicio de bitácora por fecha
    if "Fecha:" in text and any(mes in text for mes in ["junio", "mayo"]):
        current_bitacora = f"Bitácora {text}"
        print(f"\n{'='*60}")
        print(f"{current_bitacora}")
        print(f"{'='*60}")
    
    # Detectar sección de evidencias
    if "EVIDENCIAS FOTOGRÁFICAS" in text:
        print(f"\n📸 Sección de evidencias encontrada")
    
    # Buscar imágenes en el párrafo
    for run in para.runs:
        for drawing in run._element.xpath('.//a:blip'):
            embed = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if embed:
                # Obtener la ruta de la imagen
                for rel in doc.part.rels.values():
                    if rel.rId == embed:
                        image_path = rel.target_ref
                        # La imagen está en media/
                        if image_path.startswith('media/'):
                            image_name = os.path.basename(image_path)
                            source_path = os.path.join('extracted_junio/word', image_path)
                            
                            # Nombre descriptivo basado en la bitácora actual
                            safe_bitacora = current_bitacora.replace(":", "-").replace(" ", "_")[:30]
                            dest_name = f"{image_counter:02d}_{safe_bitacora}_{image_name}"
                            
                            print(f"\nImagen encontrada: {image_name}")
                            copy_image(source_path, dest_name)

# También buscar en tablas
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            
            # Detectar fecha en celda
            if "Fecha:" in text and any(mes in text for mes in ["junio", "mayo"]):
                current_bitacora = f"Bitácora_{text.replace(' ', '_')[:20]}"
                print(f"\n{'='*60}")
                print(f"{current_bitacora}")
                print(f"{'='*60}")
            
            # Detectar evidencias
            if "EVIDENCIAS FOTOGRÁFICAS" in text:
                print(f"\n📸 Sección de evidencias en tabla")
            
            # Buscar imágenes en celda
            for para in cell.paragraphs:
                for run in para.runs:
                    for drawing in run._element.xpath('.//a:blip'):
                        embed = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed:
                            for rel in doc.part.rels.values():
                                if rel.rId == embed:
                                    image_path = rel.target_ref
                                    if image_path.startswith('media/'):
                                        image_name = os.path.basename(image_path)
                                        source_path = os.path.join('extracted_junio/word', image_path)
                                        
                                        safe_bitacora = current_bitacora.replace(":", "-").replace(" ", "_")[:30]
                                        dest_name = f"{image_counter:02d}_{safe_bitacora}_{image_name}"
                                        
                                        print(f"\nImagen encontrada en tabla: {image_name}")
                                        copy_image(source_path, dest_name)

print(f"\n{'='*60}")
print(f"Total de imágenes extraídas: {image_counter}")
print(f"Ubicación: {output_dir}/")
print(f"{'='*60}")
